import json
import os
from datetime import datetime
from typing import Type, ClassVar, Optional

from langchain.tools import BaseTool
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_core.documents import Document
from pydantic import BaseModel, Field

# ============================================================
# Shared Vector Store Setup (NebulaSoft KB + optional user docs)
# ============================================================

# ✅ Must match ingest.py
NEBULA_PERSIST_DIR = "./chroma_nebula"

# ✅ Single shared embeddings instance for all tools
EMBEDDINGS = HuggingFaceEmbeddings(
    model_name="sentence-transformers/all-MiniLM-L6-v2",
    model_kwargs={"device": "cpu"},
    encode_kwargs={"normalize_embeddings": True},
)

# ✅ Permanent NebulaSoft vector store (built by ingest.py)
try:
    NEBULA_DB: Optional[Chroma] = Chroma(
        persist_directory=NEBULA_PERSIST_DIR,
        embedding_function=EMBEDDINGS,
        collection_name="nebula_kb",
    )
except Exception as e:
    print(f"⚠️ Warning: Could not load NebulaSoft vector DB: {e}")
    NEBULA_DB = None

# ✅ Temporary in-memory store for uploaded docs (will wire via UI)
TEMP_USER_DB: Optional[Chroma] = None


def load_text_from_file(file_path: str) -> str:
    """
    Load and return text from a variety of file types.
    Supports: .txt, .md, .json, .csv, .log, .py, .html, .xml
    Optionally supports: .pdf, .docx
    """

    ext = os.path.splitext(file_path)[1].lower()

    # ---- Plain Text Formats ----
    if ext in [".txt", ".md", ".log", ".py", ".html", ".xml"]:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    # ---- JSON ----
    if ext == ".json":
        import json
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
            return json.dumps(data, indent=2)

    

    # ---- PDF 
    if ext == ".pdf":
        try:
            import PyPDF2
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception:
            return "Error: PDF extraction library missing (install PyPDF2)."

    # ---- DOCX 
    if ext == ".docx":
        try:
            import docx
            document = docx.Document(file_path)
            return "\n".join([para.text for para in document.paragraphs])
        except Exception:
            return "Error: python-docx library missing."

    return f"Unsupported file type: {ext}"

def ingest_user_document(input_data: str, source_name: str = None):
    """
    Ingest either:
    - raw text string
    - or a file path to a supported text-based document
    """

    # Detect if this is a file path
    if os.path.exists(input_data):
        file_path = input_data
        text = load_text_from_file(file_path)
        source_name = source_name or os.path.basename(file_path)
    else:
        # Treat as raw text
        text = input_data
        source_name = source_name or "user_provided_text.txt"

    global TEMP_USER_DB

    doc = Document(
        page_content=text,
        metadata={
            "source": source_name,
            "kb": "user_upload",
        },
    )

    if TEMP_USER_DB is None:
        TEMP_USER_DB = Chroma.from_documents([doc], EMBEDDINGS)
    else:
        TEMP_USER_DB.add_documents([doc])

# ============================================================
# Tool 1: DocSearchTool - Searches the Vector Database
# ============================================================

class DocSearchInput(BaseModel):
    """Input schema for DocSearchTool."""
    query: str = Field(description="The search query to find relevant documentation")


class DocSearchTool(BaseTool):
    name: str = "search_documentation"
    description: str = (
        "Searches the NebulaSoft technical documentation for information about "
        "error codes, setup instructions, features, and troubleshooting. "
        "By default, it searches the official NebulaSoft manual. "
        "If additional documents have been uploaded, it also searches those. "
        "Always cite the source document in your response."
    )
    args_schema: Type[BaseModel] = DocSearchInput

    def _run(self, query: str) -> str:
        """Execute the documentation search."""
        try:
            if NEBULA_DB is None:
                return (
                    "Error: NebulaSoft documentation database is not available. "
                    "Please run ingest.py to build the vector database."
                )

            results = []

            # ✅ 1. Always search NebulaSoft base KB
            nebula_results = NEBULA_DB.similarity_search_with_score(query, k=3)
            for doc, score in nebula_results:
                results.append((doc, score, "nebula"))

            # ✅ 2. Optionally search user-uploaded docs (if any)
            if TEMP_USER_DB is not None:
                temp_results = TEMP_USER_DB.similarity_search_with_score(query, k=3)
                for doc, score in temp_results:
                    results.append((doc, score, "user"))

            if not results:
                return "No relevant documentation was found for this query."

            # ✅ 3. Sort by relevance score (ascending distance)
            results.sort(key=lambda tup: tup[1])

            # ✅ 4. Format results with clear citations
            formatted_results = []
            for doc, score, kb_type in results[:5]:
                source = doc.metadata.get("source", "Unknown")
                chunk_id = doc.metadata.get("chunk_id", "N/A")
                kb_label = "NebulaSoft Manual" if kb_type == "nebula" else "Uploaded Document"
                content = doc.page_content

                formatted_results.append(
                    f"[{kb_label} | Source: {source}, Chunk: {chunk_id}, Score: {score:.4f}]\n{content}"
                )

            return "\n\n---\n\n".join(formatted_results)

        except Exception as e:
            return f"Error searching documentation: {str(e)}"

    async def _arun(self, query: str) -> str:
        """Async version (not implemented for this project)."""
        raise NotImplementedError("Async not supported for this tool")


# ============================================================
# Tool 2: PricingCalculatorTool - Calculates Subscription Costs
# ============================================================

class PricingCalculatorInput(BaseModel):
    """Input schema for PricingCalculatorTool."""
    number_of_users: int = Field(description="Number of users for the subscription", gt=0)
    plan_type: str = Field(
        description="Type of plan: 'basic', 'pro', or 'enterprise'",
        pattern="^(basic|pro|enterprise)$",
    )


class PricingCalculatorTool(BaseTool):
    name: str = "calculate_pricing"
    description: str = (
        "Calculates the monthly subscription cost for NebulaSoft based on the number "
        "of users and plan type. Plan types are: 'basic' ($10/user), 'pro' ($20/user), "
        "or 'enterprise' ($40/user). Use this when customers ask about pricing or costs."
    )
    args_schema: Type[BaseModel] = PricingCalculatorInput

    # Pricing structure - mark as ClassVar
    PRICING: ClassVar[dict] = {
        "basic": 10,
        "pro": 20,
        "enterprise": 40,
    }

    def _run(self, number_of_users: int, plan_type: str) -> str:
        """Calculate the subscription cost."""
        plan_type = plan_type.lower()

        if plan_type not in self.PRICING:
            return f"Error: Invalid plan type '{plan_type}'. Valid options are: basic, pro, enterprise"

        price_per_user = self.PRICING[plan_type]
        total_cost = number_of_users * price_per_user

        result = (
            "Pricing Calculation:\n"
            f"Plan: {plan_type.capitalize()}\n"
            f"Number of Users: {number_of_users}\n"
            f"Price per User: ${price_per_user}/month\n"
            f"Total Monthly Cost: ${total_cost}/month"
        )

        return result

    async def _arun(self, number_of_users: int, plan_type: str) -> str:
        """Async version (not implemented for this project)."""
        raise NotImplementedError("Async not supported for this tool")


# ============================================================
# Tool 3: TicketEscalationTool - Files Support Tickets
# ============================================================

class TicketEscalationInput(BaseModel):
    """Input schema for TicketEscalationTool."""
    summary: str = Field(description="Brief summary of the issue")
    severity_level: str = Field(
        description="Severity level: 'low', 'medium', 'high'",
        pattern="^(low|medium|high)$",
    )


class TicketEscalationTool(BaseTool):
    name: str = "escalate_ticket"
    description: str = (
        "Files a support ticket for escalation to Tier-2 support. "
        "Use this tool ONLY when the documentation search cannot answer the user's question "
        "or when the issue requires human intervention. "
        "Severity levels: 'low', 'medium', 'high'."
    )
    args_schema: Type[BaseModel] = TicketEscalationInput

    def _run(self, summary: str, severity_level: str) -> str:
        """File a support ticket."""
        severity_level = severity_level.lower()

        if severity_level not in ["low", "medium", "high"]:
            return f"Error: Invalid severity level '{severity_level}'"

        # Create ticket data
        ticket = {
            "ticket_id": f"TKT-{datetime.now().strftime('%Y%m%d%H%M%S')}",
            "timestamp": datetime.now().isoformat(),
            "summary": summary,
            "severity": severity_level,
            "status": "open",
        }

        # Append to tickets.log
        try:
            with open("tickets.log", "a", encoding="utf-8") as f:
                f.write(json.dumps(ticket) + "\n")

            result = (
                "🎫 Ticket Escalated Successfully!\n"
                f"Ticket ID: {ticket['ticket_id']}\n"
                f"Severity: {severity_level.upper()}\n"
                f"Summary: {summary}\n"
                "Status: Open\n"
                "A Tier-2 support representative will contact you shortly."
            )
            return result

        except Exception as e:
            return f"Error filing ticket: {str(e)}"

    async def _arun(self, summary: str, severity_level: str) -> str:
        """Async version (not implemented for this project)."""
        raise NotImplementedError("Async not supported for this tool")

class TicketLookupInput(BaseModel):
    ticket_id: str = Field(description="The ticket ID to look up.")
    


class TicketLookupTool(BaseTool):
    name: str = "lookup_ticket"
    description: str = (
        "Looks up a support ticket by ticket ID and returns its details. "
        "Use this when the user asks about the status of an existing ticket."
    )
    args_schema: Type[BaseModel] = TicketLookupInput

    def _run(self, ticket_id: str) -> str:

        if not os.path.exists("tickets.log"):
            return "No tickets found in the system."

        try:
            with open("tickets.log", "r", encoding="utf-8") as f:
                for line in f:
                    if not line.strip():
                        continue
                    ticket = json.loads(line)
                    if ticket.get("ticket_id") == ticket_id:
                        return (
                            "🎫 Ticket Details:\n"
                            f"Ticket ID: {ticket.get('ticket_id')}\n"
                            f"Created At: {ticket.get('timestamp')}\n"
                            f"Severity: {ticket.get('severity')}\n"
                            f"Status: {ticket.get('status')}\n"
                            f"Summary: {ticket.get('summary')}"
                        )
            return f"No ticket found with ID {ticket_id}."

        except Exception as e:
            return f"Error reading ticket log: {str(e)}"

    async def _arun(self, ticket_id: str) -> str:
        raise NotImplementedError("Async not supported for this tool.")

# ============================================================
# Convenience function to get all tools
# ============================================================

def get_all_tools():
    """Returns a list of all available tools."""
    return [
        DocSearchTool(),
        PricingCalculatorTool(),
        TicketEscalationTool(),
        TicketLookupTool(),
    ]
